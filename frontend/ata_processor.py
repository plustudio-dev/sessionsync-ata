import os
import json
import re
import logging
import io
import spacy
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class AtaProcessor:
    """Classe para processar transcrições e gerar atas estruturadas."""
    
    def __init__(self, templates_dir):
        """
        Inicializa o processador de atas.
        
        Args:
            templates_dir: Diretório contendo os templates JSON de atas
        """
        self.templates_dir = templates_dir
        self.config = self._load_config()
        
        # Carregar modelo spaCy para português
        try:
            self.nlp = spacy.load("pt_core_news_sm")
            logger.info("Modelo spaCy carregado com sucesso")
        except Exception as e:
            logger.error(f"Erro ao carregar modelo spaCy: {str(e)}")
            self.nlp = None
            
        # Carregar lista de vereadores
        self.vereadores = self._load_vereadores()
        
    def _load_config(self):
        """Carrega a configuração dos templates."""
        config_path = os.path.join(self.templates_dir, 'config.json')
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Erro ao carregar configuração: {str(e)}")
            return {"tipos_sessao": {}}
    
    def _load_vereadores(self):
        """Carrega a lista de vereadores do arquivo JSON."""
        try:
            vereadores_path = os.path.join(os.path.dirname(self.templates_dir), 'static', 'vereadores.json')
            with open(vereadores_path, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return data.get('vereadores', [])
        except Exception as e:
            logger.error(f"Erro ao carregar lista de vereadores: {str(e)}")
            return []
    
    def _load_template(self, tipo_sessao):
        """
        Carrega o template para o tipo de sessão especificado.
        
        Args:
            tipo_sessao: Tipo de sessão (ordinaria, extraordinaria, solene)
            
        Returns:
            Dicionário com o template ou None se não encontrado
        """
        if tipo_sessao not in self.config['tipos_sessao']:
            logger.error(f"Tipo de sessão não encontrado: {tipo_sessao}")
            return None
            
        template_file = self.config['tipos_sessao'][tipo_sessao]
        template_path = os.path.join(self.templates_dir, template_file)
        
        try:
            with open(template_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            logger.error(f"Erro ao carregar template {template_file}: {str(e)}")
            return None
    
    def _find_section_in_transcript(self, transcript, marcadores_inicio, marcadores_fim):
        """
        Encontra uma seção na transcrição com base nos marcadores de início e fim.
        Utiliza spaCy para melhorar a identificação de seções quando disponível.
        
        Args:
            transcript: Lista de segmentos de transcrição
            marcadores_inicio: Lista de possíveis marcadores de início
            marcadores_fim: Lista de possíveis marcadores de fim
            
        Returns:
            Lista de frases encontradas na seção
        """
        section_phrases = []
        in_section = False
        
        # Concatenar todas as frases de todos os segmentos
        all_phrases = []
        for segment in transcript:
            if 'phrases' in segment:
                all_phrases.extend(segment['phrases'])
        
        # Compilar expressões regulares para os marcadores
        inicio_patterns = [re.compile(r'\b' + re.escape(marker.upper()) + r'\b') for marker in marcadores_inicio]
        fim_patterns = [re.compile(r'\b' + re.escape(marker.upper()) + r'\b') for marker in marcadores_fim]
        
        # Se temos o modelo spaCy disponível, usar NLP para melhorar a detecção
        if self.nlp:
            return self._find_section_with_nlp(all_phrases, marcadores_inicio, marcadores_fim)
        
        # Método tradicional (fallback)
        # Procurar pelos marcadores
        for i, phrase in enumerate(all_phrases):
            text = phrase.get('text', '').upper()
            
            # Verificar se encontramos um marcador de início
            if not in_section:
                for pattern in inicio_patterns:
                    if pattern.search(text):
                        in_section = True
                        section_phrases.append(phrase)
                        break
            # Se já estamos na seção, adicionar a frase
            elif in_section:
                # Verificar se encontramos um marcador de fim
                should_end = False
                for pattern in fim_patterns:
                    if pattern.search(text):
                        should_end = True
                        section_phrases.append(phrase)
                        break
                
                if should_end:
                    break
                else:
                    section_phrases.append(phrase)
        
        # Se não encontramos nenhuma frase, tentar uma abordagem mais flexível
        if not section_phrases:
            logger.info("Tentando encontrar seção com correspondência parcial...")
            in_section = False
            
            for phrase in all_phrases:
                text = phrase.get('text', '').upper()
                
                # Verificar se encontramos um marcador de início (correspondência parcial)
                if not in_section:
                    for marker in marcadores_inicio:
                        if marker.upper() in text:
                            in_section = True
                            section_phrases.append(phrase)
                            break
                # Se já estamos na seção, adicionar a frase
                elif in_section:
                    # Verificar se encontramos um marcador de fim
                    should_end = False
                    for marker in marcadores_fim:
                        if marker.upper() in text:
                            should_end = True
                            break
                    
                    if should_end:
                        section_phrases.append(phrase)
                        break
                    else:
                        section_phrases.append(phrase)
        
        return section_phrases
        
    def _find_section_with_nlp(self, phrases, marcadores_inicio, marcadores_fim):
        """
        Encontra uma seção na transcrição usando processamento de linguagem natural.
        
        Args:
            phrases: Lista de frases da transcrição
            marcadores_inicio: Lista de possíveis marcadores de início
            marcadores_fim: Lista de possíveis marcadores de fim
            
        Returns:
            Lista de frases encontradas na seção
        """
        section_phrases = []
        in_section = False
        
        # Preparar marcadores para comparação semântica
        inicio_docs = [self.nlp(marker.lower()) for marker in marcadores_inicio]
        fim_docs = [self.nlp(marker.lower()) for marker in marcadores_fim]
        
        for phrase in phrases:
            if 'text' not in phrase:
                continue
                
            text = phrase.get('text', '')
            doc = self.nlp(text.lower())
            
            # Verificar se encontramos um marcador de início
            if not in_section:
                # Primeiro, verificar correspondência exata (mais confiável)
                for pattern in [re.compile(r'\b' + re.escape(marker.upper()) + r'\b') for marker in marcadores_inicio]:
                    if pattern.search(text.upper()):
                        in_section = True
                        section_phrases.append(phrase)
                        logger.info(f"Início de seção encontrado (regex): {text}")
                        break
                        
                # Se não encontrou por regex, tentar similaridade semântica
                if not in_section:
                    for i, inicio_doc in enumerate(inicio_docs):
                        # Verificar similaridade com frases curtas
                        if len(doc) < 20 and doc.similarity(inicio_doc) > 0.75:
                            in_section = True
                            section_phrases.append(phrase)
                            logger.info(f"Início de seção encontrado (semântica): {text} -> {marcadores_inicio[i]}")
                            break
                        # Para frases longas, verificar se contém o marcador
                        elif any(token.text.lower() in marcadores_inicio[i].lower() for token in doc):
                            in_section = True
                            section_phrases.append(phrase)
                            logger.info(f"Início de seção encontrado (token): {text} -> {marcadores_inicio[i]}")
                            break
            
            # Se já estamos na seção, adicionar a frase
            elif in_section:
                # Verificar se encontramos um marcador de fim
                should_end = False
                
                # Primeiro, verificar correspondência exata (mais confiável)
                for pattern in [re.compile(r'\b' + re.escape(marker.upper()) + r'\b') for marker in marcadores_fim]:
                    if pattern.search(text.upper()):
                        should_end = True
                        section_phrases.append(phrase)
                        logger.info(f"Fim de seção encontrado (regex): {text}")
                        break
                
                # Se não encontrou por regex, tentar similaridade semântica
                if not should_end:
                    for i, fim_doc in enumerate(fim_docs):
                        # Verificar similaridade com frases curtas
                        if len(doc) < 20 and doc.similarity(fim_doc) > 0.75:
                            should_end = True
                            section_phrases.append(phrase)
                            logger.info(f"Fim de seção encontrado (semântica): {text} -> {marcadores_fim[i]}")
                            break
                        # Para frases longas, verificar se contém o marcador
                        elif any(token.text.lower() in marcadores_fim[i].lower() for token in doc):
                            should_end = True
                            section_phrases.append(phrase)
                            logger.info(f"Fim de seção encontrado (token): {text} -> {marcadores_fim[i]}")
                            break
                
                if should_end:
                    break
                else:
                    section_phrases.append(phrase)
        
        return section_phrases
    
    def _format_section_text(self, phrases):
        """
        Formata o texto de uma seção a partir das frases encontradas.
        Utiliza spaCy para melhorar a formatação e identificação de entidades quando disponível.
        
        Args:
            phrases: Lista de frases da seção
            
        Returns:
            Texto formatado da seção
        """
        if not phrases:
            return ""
            
        text_parts = []
        current_speaker = None
        
        # Verificar se podemos usar NLP para melhorar a formatação
        if self.nlp and self.vereadores:
            return self._format_section_with_nlp(phrases)
        
        # Método tradicional (fallback)
        for phrase in phrases:
            if 'text' not in phrase:
                continue
                
            text = phrase['text'].strip()
            speaker = phrase.get('speaker', None)
            
            # Adicionar identificação do orador se disponível e diferente do anterior
            if speaker and speaker != current_speaker:
                text_parts.append(f"\n{speaker.upper()}: {text}")
                current_speaker = speaker
            else:
                # Se for o mesmo orador ou não houver identificação, apenas adicionar o texto
                text_parts.append(text)
        
        # Juntar as partes com espaços, mas preservar quebras de linha
        result = " ".join(text_parts)
        
        # Normalizar espaços duplos e quebras de linha
        result = re.sub(r' +', ' ', result)
        result = re.sub(r'\n +', '\n', result)
        
        return result
        
    def _format_section_with_nlp(self, phrases):
        """
        Formata o texto de uma seção usando processamento de linguagem natural.
        Melhora a identificação de oradores e entidades nomeadas.
        
        Args:
            phrases: Lista de frases da seção
            
        Returns:
            Texto formatado da seção
        """
        if not phrases:
            return ""
            
        text_parts = []
        current_speaker = None
        
        # Pré-processar os nomes dos vereadores para facilitar a comparação
        vereadores_docs = [self.nlp(vereador.lower()) for vereador in self.vereadores]
        vereadores_tokens = [set(doc.text.lower().split()) for doc in vereadores_docs]
        
        for phrase in phrases:
            if 'text' not in phrase:
                continue
                
            text = phrase['text'].strip()
            speaker = phrase.get('speaker', None)
            
            # Processar o texto com spaCy
            doc = self.nlp(text)
            
            # Tentar identificar o orador se não estiver explícito
            identified_speaker = None
            
            # Se já temos um orador identificado, usar ele
            if speaker:
                identified_speaker = speaker
            # Caso contrário, tentar identificar o orador no texto
            else:
                # Verificar entidades de pessoa
                person_entities = [ent.text for ent in doc.ents if ent.label_ == 'PER']
                
                # Se encontramos entidades de pessoa, verificar se alguma corresponde a um vereador
                if person_entities:
                    for entity in person_entities:
                        entity_doc = self.nlp(entity.lower())
                        entity_tokens = set(entity_doc.text.lower().split())
                        
                        # Verificar correspondência com vereadores
                        for i, vereador_tokens in enumerate(vereadores_tokens):
                            # Se há interseção significativa entre os tokens
                            if len(entity_tokens.intersection(vereador_tokens)) >= 2:
                                identified_speaker = self.vereadores[i]
                                break
                        
                        if identified_speaker:
                            break
            
            # Adicionar identificação do orador se disponível e diferente do anterior
            if identified_speaker and identified_speaker != current_speaker:
                text_parts.append(f"\n{identified_speaker.upper()}: {text}")
                current_speaker = identified_speaker
            else:
                # Verificar se o texto começa com um nome de vereador
                speaker_found = False
                doc_tokens = set(doc.text.lower().split()[:3])  # Primeiras 3 palavras
                
                for i, vereador_tokens in enumerate(vereadores_tokens):
                    # Se há interseção significativa entre os tokens iniciais
                    if len(doc_tokens.intersection(vereador_tokens)) >= 2:
                        # Extrair o nome do vereador do início e o restante do texto
                        vereador = self.vereadores[i]
                        # Encontrar onde termina o nome do vereador no texto
                        vereador_parts = vereador.lower().split()
                        text_parts_lower = text.lower().split()
                        
                        # Encontrar o índice onde termina o nome do vereador
                        end_idx = 0
                        for part in vereador_parts:
                            if part in text_parts_lower[end_idx:end_idx+3]:
                                end_idx = text_parts_lower.index(part, end_idx) + 1
                        
                        # Separar o texto
                        remaining_text = text.split(' ', end_idx)[end_idx:]
                        remaining_text = ' '.join(remaining_text)
                        
                        text_parts.append(f"\n{vereador.upper()}: {remaining_text}")
                        current_speaker = vereador
                        speaker_found = True
                        break
                
                if not speaker_found:
                    # Se for o mesmo orador ou não houver identificação, apenas adicionar o texto
                    text_parts.append(text)
        
        # Juntar as partes com espaços, mas preservar quebras de linha
        result = " ".join(text_parts)
        
        # Normalizar espaços duplos e quebras de linha
        result = re.sub(r' +', ' ', result)
        result = re.sub(r'\n +', '\n', result)
        
        # Melhorar a formatação de entidades importantes
        if self.nlp:
            result = self._enhance_entities(result)
        
        return result
        
    def _enhance_entities(self, text):
        """
        Melhora a formatação de entidades importantes no texto.
        
        Args:
            text: Texto a ser processado
            
        Returns:
            Texto com entidades formatadas
        """
        doc = self.nlp(text)
        enhanced_text = text
        
        # Identificar e formatar entidades importantes
        for ent in doc.ents:
            if ent.label_ in ['LOC', 'ORG', 'PER']:
                # Formatar entidades em maiúsculas
                enhanced_text = enhanced_text.replace(ent.text, ent.text.upper())
            elif ent.label_ in ['MISC', 'LAW']:
                # Formatar referências a leis e documentos em itálico (para HTML)
                if '<i>' not in ent.text and '</i>' not in ent.text:
                    enhanced_text = enhanced_text.replace(ent.text, f"<i>{ent.text}</i>")
        
        # Identificar e formatar números de projetos de lei
        projeto_pattern = re.compile(r'\b(PROJETO|PROJETOS)\s+DE\s+(LEI|DECRETO)\s+N[º°]\s*\d+\b', re.IGNORECASE)
        for match in projeto_pattern.finditer(text):
            projeto = match.group(0)
            enhanced_text = enhanced_text.replace(projeto, f"<b>{projeto.upper()}</b>")
        
        # Identificar e formatar votações
        votacao_pattern = re.compile(r'\b(APROVAD[OA]|REJEITAD[OA])\s+(POR\s+)?(UNANIMIDADE|MAIORIA)\b', re.IGNORECASE)
        for match in votacao_pattern.finditer(text):
            votacao = match.group(0)
            enhanced_text = enhanced_text.replace(votacao, f"<b>{votacao.upper()}</b>")
        
        return enhanced_text
    
    def _apply_template_variables(self, template_text, variables):
        """
        Aplica as variáveis ao template.
        
        Args:
            template_text: Texto do template com placeholders
            variables: Dicionário de variáveis para substituir
            
        Returns:
            Texto com as variáveis substituídas
        """
        result = template_text
        for key, value in variables.items():
            placeholder = "{" + key + "}"
            result = result.replace(placeholder, str(value))
        return result
    
    def process_transcript(self, transcript_data, tipo_sessao, metadata):
        """
        Processa a transcrição e gera uma ata estruturada.
        
        Args:
            transcript_data: Dados da transcrição (lista de segmentos)
            tipo_sessao: Tipo de sessão (ordinaria, extraordinaria, solene)
            metadata: Metadados adicionais para preencher o template
            
        Returns:
            Texto da ata estruturada ou None se ocorrer erro
        """
        template = self._load_template(tipo_sessao)
        if not template:
            return None
            
        ata_parts = []
        
        # Processar cabeçalho
        if 'cabecalho' in template:
            cabecalho = template['cabecalho']
            
            # Estrutura
            if 'estrutura' in cabecalho:
                ata_parts.append(self._apply_template_variables(cabecalho['estrutura'], metadata))
                
            # Presidência
            if 'presidencia' in cabecalho:
                ata_parts.append(self._apply_template_variables(cabecalho['presidencia'], metadata))
                
            # Presença
            if 'presenca' in cabecalho:
                ata_parts.append(self._apply_template_variables(cabecalho['presenca'], metadata))
        
        # Processar corpo
        if 'corpo' in template:
            corpo = template['corpo']
            
            # Abertura
            if 'abertura' in corpo:
                for linha in corpo['abertura']:
                    ata_parts.append(self._apply_template_variables(linha, metadata))
            
            # Expediente (apenas para sessões ordinárias)
            if tipo_sessao == 'ordinaria' and 'expediente' in corpo:
                expediente = corpo['expediente']
                if 'marcadores' in expediente:
                    expediente_phrases = self._find_section_in_transcript(
                        transcript_data,
                        expediente['marcadores']['inicio'],
                        expediente['marcadores']['fim']
                    )
                    
                    if expediente_phrases:
                        if 'instrucao' in expediente:
                            ata_parts.append(expediente['instrucao'])
                        ata_parts.append(self._format_section_text(expediente_phrases))
            
            # Pronunciamentos
            if 'pronunciamentos' in corpo:
                pronunciamentos = corpo['pronunciamentos']
                if 'marcadores' in pronunciamentos:
                    pronunciamento_phrases = self._find_section_in_transcript(
                        transcript_data,
                        pronunciamentos['marcadores']['inicio'],
                        pronunciamentos['marcadores']['fim']
                    )
                    
                    if pronunciamento_phrases:
                        if 'instrucao' in pronunciamentos:
                            ata_parts.append(pronunciamentos['instrucao'])
                        ata_parts.append(self._format_section_text(pronunciamento_phrases))
            
            # Ordem do dia
            if 'ordem_do_dia' in corpo:
                ordem = corpo['ordem_do_dia']
                if 'marcadores' in ordem:
                    ordem_phrases = self._find_section_in_transcript(
                        transcript_data,
                        ordem['marcadores']['inicio'],
                        ordem['marcadores']['fim']
                    )
                    
                    if ordem_phrases:
                        if 'instrucao' in ordem:
                            ata_parts.append(ordem['instrucao'])
                        ata_parts.append(self._format_section_text(ordem_phrases))
            
            # Homenagens (apenas para sessões solenes)
            if tipo_sessao == 'solene' and 'homenagens' in corpo:
                homenagens = corpo['homenagens']
                if 'marcadores' in homenagens:
                    homenagens_phrases = self._find_section_in_transcript(
                        transcript_data,
                        homenagens['marcadores']['inicio'],
                        homenagens['marcadores']['fim']
                    )
                    
                    if homenagens_phrases:
                        if 'instrucao' in homenagens:
                            ata_parts.append(homenagens['instrucao'])
                        ata_parts.append(self._format_section_text(homenagens_phrases))
        
        # Processar encerramento
        if 'encerramento' in template:
            encerramento = template['encerramento']
            
            if 'marcadores' in encerramento:
                encerramento_phrases = self._find_section_in_transcript(
                    transcript_data,
                    encerramento['marcadores']['inicio'],
                    encerramento['marcadores']['fim']
                )
                
                if encerramento_phrases:
                    ata_parts.append(self._format_section_text(encerramento_phrases))
                elif 'template' in encerramento:
                    ata_parts.append(self._apply_template_variables(encerramento['template'], metadata))
        
        # Juntar todas as partes com quebras de linha duplas
        return "\n\n".join(ata_parts)
    
    def generate_docx(self, ata_text, metadata):
        """
        Gera um documento Word a partir do texto da ata.
        
        Args:
            ata_text: Texto da ata estruturada
            metadata: Metadados da sessão
            
        Returns:
            Objeto BytesIO contendo o documento Word
        """
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Adicionar título
        title = doc.add_heading(level=1)
        title_run = title.add_run("ATA DA SESSÃO")
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar metadados
        if 'title' in metadata and metadata['title']:
            subtitle = doc.add_paragraph()
            subtitle_run = subtitle.add_run(metadata['title'])
            subtitle_run.font.size = Pt(14)
            subtitle_run.font.bold = True
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if 'date' in metadata and metadata['date']:
            date_para = doc.add_paragraph()
            date_run = date_para.add_run(f"Data: {metadata['date']}")
            date_run.font.size = Pt(12)
            date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Adicionar linha horizontal
        doc.add_paragraph("_" * 70)
        
        # Adicionar conteúdo da ata
        for paragraph_text in ata_text.split("\n\n"):
            if paragraph_text.strip():
                p = doc.add_paragraph()
                p_run = p.add_run(paragraph_text.strip())
                p_run.font.size = Pt(12)
        
        # Salvar o documento em memória
        docx_bytes = io.BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes
